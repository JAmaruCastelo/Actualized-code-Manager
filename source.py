

# PAQUETES A IMPORTAR
# importamos
# pip install python-docx
# pip install scikit-bio
from matplotlib.patches import Polygon
import random
import numpy as np
import matplotlib.pyplot as plt
import ecopy as ep
from collections import Counter
from docx import Document
from datetime import datetime
import pandas as pd
from scipy.special import comb
from pandas import DataFrame
import matplotlib.gridspec as gsp

def crear_lista_de_cantidades_de_veces(df_new2):
    """ funcion que becesita el df con el form de epicollet para obtener una lista 
    con las cantidad de cada especie uno por uno"""
    lista = []
    for indice in range(len(df_new2)):
        cantidad = df_new2["2_INDIVIDUOS"][indice]
        especie = df_new2["1_ESPECIE"][indice]
        for indice in range(cantidad):
            lista.append(especie)
    return lista


def añadir_lista_to_df(df, add, name):
    """ esta es una funcion que permite añadir listas a un df
    se necesita un df y una lista para añadir, y un nombre
    devuelve un df con una fila adicional con la lista y el nombre
    """
    new_df = df[name] = add
    return new_df

def diccionario(name_routes, name1="trail", name2="id"):
    "necesita un dtaframe, y los noimbres de dos filas y devuelve un dicccionario con las 2 filas"
    return dict(zip(name_routes[name1], name_routes[name2]))

diccio_name=[]
diccio_name2=[]
def crear_base_datos(df_new1, diccio=diccio_name, diccio2=diccio_name2):
    """nota: volver a definir para evitar el error de diccio y diccio_name2"""
    dataframe_new = pd.DataFrame()
    Caja = df_new1["5_CAJA"]
    Pais = ["Perú"]*len(df_new1)
    Departamento = ["Madre de Dios"]*len(df_new1)
    provincia = ["Manu"]*len(df_new1)
    Distrito = ["Salvación"]*len(df_new1)
    Localidad = ["Salvación"]*len(df_new1)
    punto_de_colecta = df_new1["4_CODIGO"]
    Modo = df_new1["9_TIPO_DE_MUESTREO"]
    # las fechas deben estar en formato de fecha
    Fecha = df_new1["7_FECHA_DE_COLECTA"]
    Tipo = df_new1["8_TIPO_DE_BOSQUE"]
    Fecha2 = []
    for e in Fecha:
        e = str(e)
        try:
            date = datetime.strptime(e, '%d/%m/%Y')
            Fecha2.append(date)
        except ValueError:
            Fecha2.append("none")
    Latitud = ["-12.789"]*len(df_new1)
    Longitud = ["-71.398"]*len(df_new1)
    Altitud = ["496msnm"]*len(df_new1)
    Colector = ["CREES"]*len(df_new1)
    Codigo = df_new1["ec5_branch_owner_uuid"]
    Identificacion = []
    for e in Codigo:
        nuevo = diccio2[e]
        Identificacion.append(nuevo)
    Autor = []
    for e in Identificacion:
        try:
            autor = diccio[e]
        except KeyError:
            autor = "nofound"
        Autor.append(autor)
    Identificador = ["Amaru,J"]*len(df_new1)
    dataframe_new["Caja"] = Caja
    dataframe_new["Pais"] = Pais
    dataframe_new["Departamento"] = Departamento
    dataframe_new["provincia"] = provincia
    dataframe_new["Distrito"] = Distrito
    dataframe_new["Localidad"] = Localidad
    dataframe_new["punto_de_colecta"] = punto_de_colecta
    dataframe_new["Modo"] = Modo
    dataframe_new["Fecha"] = Fecha2
    dataframe_new["Latitud"] = Latitud
    dataframe_new["Longitud"] = Longitud
    dataframe_new["Altitud"] = Altitud
    dataframe_new["Colector"] = Colector
    dataframe_new["Identificacion"] = Identificacion
    dataframe_new["Autor"] = Autor
    dataframe_new["Identificador"] = Identificador
    dataframe_new["Tipo"] = Tipo
    return dataframe_new


def años(df, Fecha="Fecha"):
    """ esta funcion te da una lista de los años que se puso en la fecha
    necesita un dataframe que tenga una columna denominada Fecha con fechas
    devuelve una lista con los años, si no es un año devuelve "none"""
    años = []
    for elemento in range(0, len(df)):
        fecha = df["Fecha"][elemento]
        if isinstance(fecha, datetime):
            a = fecha.year
            años.append(a)
        elif isinstance(fecha, int):
            a = "none"
            años.append(a)
        else:
            a = "none"
            años.append(a)
    return años

def seleccionar(dfA, año):
    """ funcion para seleccionar de el dataframne un alño requerido
    se requiere un data frame y el año
    y devuelve un nuevo dataframe con los años requeridos"""
    return dfA[dfA["años"] == año]


def seleccionar_caracteristica(df, requiere, fila):
    """ funcion para seleccionar de el dataframne una caracteristica que se requiere 
    de la fila dada se debe poner el nombre o la caracteristica que se requiere
    y devuelve un nuevo dataframe de solo los datos requeridos similar a la funcion
    seleccionar de años"""
    return df[df[fila] == requiere]


def añadir_tipo_bosque(df):
    TIPO = []
    for elemento in range(0, len(df)):
        pc = df["punto_de_colecta"][elemento]
        if isinstance(pc, str):
            if len(pc) > 3:
                Tipo = pc[0:3]
                TIPO.append(Tipo)
            else:
                TIPO.append("none")
        else:
            TIPO.append("none")
    df["Tipo"] = TIPO

def completar(dfB, diccionarito, code="trail", newname="ide"):
    """ este dataframe modifica un dataframe añadiendole una nueva fila con el resultadoi de un
    diccionario requiere, un df y un diccionariuo, tambien se le puede colocar opcionalmente el nombre de 
    la dila de el datafrme que se buscara en el diccionario y el nuevo nombre de la fila"""
    a = []
    trails = dfB[code]
    for trail in trails:
        try:
            ide = diccionarito[trail]
            a.append(ide)
        except KeyError:
            a.append("nofound")
    dfB[newname] = a

def obtener_codigo_añadido(dfA, name1="punto_de_colecta", name2="Modo", name3="Fecha"):
    """ el codigo es con la primera letra de el name1, la primera letra de el name 2 
    y todo el name 3. los names son los nombres de las columnas de la filas de el dataframe"""
    data = []
    for indice in range(0, len(dfA)):
        primera = str(dfA.loc[indice][name1])  # puntos de colecta
        segunda = str(dfA.loc[indice][name2])
        tercera = str(dfA.loc[indice][name3])
        datito = primera[0]+segunda[0]+tercera
        data.append(datito)
    añadir_lista_to_df(dfA, data, name="codigo-dicti")


def crear_etiquetas(dfA):
    etiquetas_word = Document()

    for indice in range(0, len(dfA)):
        pais = dfA.loc[indice]["Pais"]
        departamento = dfA.loc[indice]["Departamento"]
        provincia = dfA.loc[indice]["provincia"]
        distrito = dfA.loc[indice]["Distrito"]
        localidad = dfA.loc[indice]["Localidad"]
        point = dfA.loc[indice]["punto_de_colecta"]
        modo = dfA.loc[indice]["Modo"]
        trail = dfA.loc[indice]["trail"]
        Tipo = dfA.loc[indice]["Tipo"]
        # latitud=str(dfA.loc[indice]["latitud"])
        # longitud=str(dfA.loc[indice]["longitud"])
        # latitud=str(dfA.loc[indice]["Altitud"])
        Fecha = dfA.loc[indice]["Fecha"]
        if isinstance(Fecha, datetime):
            fecha = Fecha.date()
        else:
            fecha = ""
        Colector = dfA.loc[indice]["Colector"]

        etiquetas_word.add_paragraph(f"{pais.upper()}: {departamento}")
        etiquetas_word.add_paragraph(
            f"{provincia}, {distrito[0:len(distrito)-2]}, {localidad}")
        etiquetas_word.add_paragraph(f"{point}  {modo} {trail}")
        #print(f"{latitud[0:4]}  {longitud[0:4]} {latitud[0:4]}")
        etiquetas_word.add_paragraph(f"-12.789 -71.398 496msnm")
        etiquetas_word.add_paragraph(f"{fecha} {Colector} {Tipo}")

    return etiquetas_word.save('etiquetas.docx')

def unir_df(DF1, DF2, DF3, DF4, DF5, DF6):
    """ esta funcion te permite unir 4 dataframes obtenidos de eexcel y epicollet
    necesita  las rutas de los siguientes datafrmaes :df1 que es un fatafrmae de excel, df2 que es un dataframe de epicolet al igual que los otrs ya 
    arreglados con las anteiores funciones devuelve el dataframe unido"""
    df1 = pd.read_excel(DF1)
    añadir_tipo_bosque(df1)
    df2 = pd.read_excel(DF2)
    df2 = df2.drop(['Unnamed: 0'], axis=1)
    df3 = pd.read_excel(DF3)
    df3 = df3.drop(['Unnamed: 0'], axis=1)
    df4 = pd.read_excel(DF4)
    df4 = df4.drop(['Unnamed: 0'], axis=1)
    df5 = pd.read_excel(DF5)
    df5 = df5.drop(['Unnamed: 0'], axis=1)
    df6 = pd.read_excel(DF6)
    df6 = df6.drop(['Unnamed: 0'], axis=1)
    final = pd.concat([df1, df2, df3, df4, df5, df6])
    return final

# esto es temporal es una modificacion para poder trabajar con los datos de familias


def unir_df(DF1, DF2):
    """ esta funcion te permite unir 4 dataframes obtenidos de eexcel y epicollet
    necesita  las rutas de los siguientes datafrmaes :df1 que es un fatafrmae de excel, df2 que es un dataframe de epicolet al igual que los otrs ya 
    arreglados con las anteiores funciones devuelve el dataframe unido"""
    df1 = pd.read_excel(DF1)
    añadir_tipo_bosque(df1)
    df2 = pd.read_excel(DF2)
    df2 = df2.drop(['Unnamed: 0'], axis=1)
    final = pd.concat([df1, df2])
    return final

############################################################################

def tipo_bosque(base, criterio="Identificacion", criterio2="Tipo", criterio3=["CCR","PCR","SLR"]):
    """ funcion que te permite tener un data frame de especies por lugares de muestreo
    se necesita un dataframe que tenga las especies y el tipo de bosque de cada especies"""
    DF1 = pd.DataFrame()
    CCR = Counter(base[base[criterio2] == criterio3[0]][criterio])
    PCR = Counter(base[base[criterio2] == criterio3[1]][criterio])
    SLR = Counter(base[base[criterio2] == criterio3[2]][criterio])
    DF1["ESPECIE"] = base[criterio].unique()

    lista_CCR = []
    for e in DF1["ESPECIE"]:
        lista_CCR.append(CCR[e])
    DF1[criterio3[0]] = lista_CCR

    lista_PCR = []
    for e in DF1["ESPECIE"]:
        lista_PCR.append(PCR[e])
    DF1[criterio3[1]] = lista_PCR

    lista_SLR = []
    for e in DF1["ESPECIE"]:
        lista_SLR.append(SLR[e])
    DF1[ criterio3[2]] = lista_SLR
    DF1.index = base[criterio].unique()
    DF2 = DF1.drop(['ESPECIE'], axis=1)
    return DF2


def tipo_meses(base, criterio="Identificacion"):
    """ funcion que te permite tener un data frame de especies por lugares de muestreo
    se necesita un dataframe que tenga las especies y el tipo de bosque de cada especies"""
    DF1 = pd.DataFrame()
    DF1["ESPECIE"] = base[criterio].unique()
    for e in range(1, 13):
        numero = Counter(base[base["mes"] == e][criterio])
        lista_numero = []
        for e2 in DF1["ESPECIE"]:
            lista_numero.append(numero[e2])
        DF1[f"mes: {e}"] = lista_numero
    return DF1


def tipo_trap(base, criterio="Identificacion"):
    """ funcion que te permite tener un data frame de especies por lugares de muestreo
    se necesita un dataframe que tenga las especies y el tipo de bosque de cada especies"""
    DF1 = pd.DataFrame()
    DF1["ESPECIE"] = base[criterio].unique()
    trap = ["BAM", "PF", "PF-CEBO", "MAL"]
    for e in trap:
        numero = Counter(base[base["tramp"] == e][criterio])
        lista_numero = []
        for e2 in DF1["ESPECIE"]:
            lista_numero.append(numero[e2])
        DF1[f"{e}"] = lista_numero
    return DF1

def tipo_bosque_idem(base, criterio="Identificacion", criterio2="tipoide"):
    DF1 = pd.DataFrame()
    DF1["ESPECIE"] = base[criterio].unique()
    trap = base[criterio2].unique()
    for e in trap:
        numero = Counter(base[base[criterio2] == e][criterio])
        lista_numero = []
        for e2 in DF1["ESPECIE"]:
            lista_numero.append(numero[e2])
        DF1[f"{e}"] = lista_numero
    return DF1

def tipo_meses_tramp(base, criterio="Identificacion"):
    """ funcion que te permite tener un data frame de especies por lugares de muestreo
    se necesita un dataframe que tenga las especies y el tipo de bosque de cada especies"""
    tipo_mes_tramp = base["tramp"]+[str(e) for e in base["mes"]]
    base["mesxtramp"] = tipo_mes_tramp
    DF1 = pd.DataFrame()
    DF1["ESPECIE"] = base[criterio].unique()
    a = tipo_mes_tramp.unique()
    for e in a:
        numero = Counter(base[base["mesxtramp"] == e][criterio])
        lista_numero = []
        for e2 in DF1["ESPECIE"]:
            lista_numero.append(numero[e2])
        DF1[f"mes: {e}"] = lista_numero
    return DF1

def graficos_seleccionados(lista, por_tipo, estilo=["ob--", "xr--"]):
    """necesita una lista de nombres de especies que nos interesan y devuelve y devuelve
    el cuadro por tipo de las especies seleccionadas"""
    ax = plt.axes()
    for e in lista:
        df2 = por_tipo[por_tipo.index == e]
        a = random.choice(estilo)
        plt.plot(df2.transpose(), a)
        estilo.remove(a)
        plt.text("SLR", df2["SLR"], e, fontfamily="Helvetica", fontstyle="italic")
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    plt.ylabel("ABUNDANCIA", fontfamily='Helvetica', fontsize=10)
    
def grafica_r_a(a, name=None, ticks=True, scale=False):
    """ realiza la curvas de rango abundancia  necesita una serie proveniente de
    el data_frame_solo que ya se encuentra ordenado"""
    a=a.sort_values(ascending=False)
    indice = a.index.values
    indice2=a.index.values
    ax = plt.axes()
    plt.plot(indice, a, "ob--")
    plt.xticks([])
    plt.xticks(ticks=np.arange(0, len(a)+4, 1),labels=[],
               rotation=45, fontfamily="Helvetica", fontstyle="italic", 
               fontsize=10)
    plt.ylabel('Log10 (n/N)')
    z=1
    for nombre in indice2 :
        nombre=str(nombre)
        yubication = a[nombre]
        plt.annotate(nombre,(z-1, yubication), rotation=45,fontsize=7)
        z=z+1
    plt.title(name)
    if scale==True:
        plt.yscale("log")
    plt.yticks(ticks=np.arange(int(min(a))-0.5, int(max(a))+1, 0.5),labels=np.arange(int(min(a))-0.5, int(max(a))+1, 0.5))
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)


def data_frame_solo(por_tipo, tipos=["CCR", "PCR", "SLR"]):
    b = []
    for e in tipos:
        a = por_tipo[e]
        a = a.sort_values(ascending=False)
        a = a[a != 0]
        b.append(a)
    return b
### prueba beta
def unir_general(por_tipo, tipos=["CCR", "PCR", "SLR"]):
    """une diferentes filas de un df para realizar grafica rango-A"""
    b=data_frame_solo(por_tipo, tipos=tipos)
    l=len(b)
    unido=pd.DataFrame()
    for e in range (0,l):
        data=b[e]
        data=pd.DataFrame(data)
        dado=[]
        for e in range(len(data)):
            dado.append(data.columns[0][0])
        data["t"]=data.index+"_"+dado
        data=data.set_index("t")
        data.columns=["general"]
        unido=pd.concat([unido,data], axis = 0)
    return unido
#### fin prueba beta

##### base de datos generalizado
def grafi_general_r_a (por_tipo, tipos=["CCR", "PCR", "SLR"]):
    """funcion que recibe una base de datos como en por_tipo de 3 filas
    que te devuelve un grafico de rango abundancia de las 3 unidas"""
    data=data_frame_solo(por_tipo, tipos=tipos)
    a=data[0]
    a=a/sum(a)
    a=np.log10(a)    
    b=data[1]
    b=b/sum(b)
    b=np.log10(b)
    c=data[2]
    c=c/sum(c)
    c=np.log10(c)
    
    fig = plt.figure(figsize=(12,4))
    
    g=gsp.GridSpec(1,3, figure=fig)
    ax1 = fig.add_subplot(g[0,0:1])
    ax1.patch.set_alpha(0)
    ax2 = fig.add_subplot(g[0,1:2], sharey=ax1)
    ax2.patch.set_alpha(0)
    ax3 = fig.add_subplot(g[0,2:3], sharey=ax1)
    ax3.patch.set_alpha(0)

    #### ponemos limites de el eje x
    xmax=max(len(a), len(b), len(c))
    #### ponemos limites de el eje y
    ymin=int(max(max(a),max(b),max(c)))
    ymax=int(min(min(a), min(b), min(c)))
    #### primer eje
    indice = a.index.values
    indice2= a.index.values
    ax1.tick_params(axis='y', which='both', bottom=False, right=False, top=False) 
    ax1.tick_params(axis='x', which='both', right=False,left=False,bottom=False,labelbottom=False) 
    ax1.spines['right'].set_visible(False)
    ax1.spines['top'].set_visible(False)
    ax1.plot(indice,a , color="black")
    ax1.set_xticks(ticks=np.arange(0, xmax, 1),labels=[],rotation=45, fontfamily="Helvetica", fontstyle="italic", fontsize=25)
    ax1.set_ylabel("Log10 (n/N)",fontsize=15)
    ax1.set_xlabel(tipos[0],fontsize=15)
    z=1
    for nombre in indice2 :
        nombre=str(nombre)
        yubication = a[nombre]
        ax1.annotate(nombre,(z-1, yubication+0.05), rotation=45,fontsize=12)
        z=z+1
    ax1.set_yticks(ticks=np.arange(int(min(a))-0.5, int(max(a))+1, 0.5),labels=np.arange(int(min(a))-0.5, int(max(a))+1, 0.5))

    #### segundo eje
    indice = b.index.values
    indice2= b.index.values    
    ax2.tick_params(axis='y', which='both', bottom=False, left=False, right=False, top=False, labelleft=False)
    ax2.tick_params(axis='x', which='both', bottom=False,left=False, right=False,top=False,labelbottom=False) 
    ax2.spines['left'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['top'].set_visible(False)
    ax2.plot(indice,b , color="black")
    ax2.set_xticks(ticks=np.arange(0, xmax, 1),labels=[], rotation=45, fontfamily="Helvetica", fontstyle="italic", fontsize=25)
    ax2.set_xlabel(tipos[1],fontsize=15)
    z=1
    for nombre in indice2 :
        nombre=str(nombre)
        yubication = b[nombre]
        ax2.annotate(nombre,(z-1, yubication+0.05), rotation=45,fontsize=12)
        z=z+1
    #ax2.set_yticks(ticks=np.arange(int(min(b))-0.5, int(max(b))+1, 0.5),labels=np.arange(int(min(b))-0.5, int(max(b))+1, 0.5))

    #### tercer eje
    indice = c.index.values
    indice2= c.index.values 
    ax3.tick_params(axis='y', which='both', bottom=False, left=False, right=False, top=False,labelleft=False) 
    ax3.tick_params(axis='x', which='both', bottom=False,left=False, right=False, labelbottom=False) 
    ax3.spines['left'].set_visible(False)
    ax3.spines['right'].set_visible(False)
    ax3.spines['top'].set_visible(False)
    ax3.plot(indice,c, color="black")
    ax3.set_xticks(ticks=np.arange(0, xmax, 1),labels=[],rotation=45, fontfamily="Helvetica", fontstyle="italic", fontsize=25)
    ax3.set_xlabel(tipos[2],fontsize=15)
    z=1
    for nombre in indice2 :
        nombre=str(nombre)
        yubication = c[nombre]
        ax3.annotate(nombre,(z-1, yubication+0.05), rotation=45,fontsize=12)
        z=z+1
    return fig
     

def guardar_grafricos (data_frame_solo, names=["CCR","PCR", "SLR"]):
    """funcion para guardar graficos de rango abundancia automaticamente
    requiere, la lista de dataframe obtenidas de la funcion data_frame_solo
    se guardan las imagenes en donde se encuentren los archivos
    """
    for e in range (len (data_frame_solo)):
        a=b[e]
        plot=grafica_r_a(a, name=names[e])
        plot.savefig(f"{names[e]}_rang_abundance.jpg")
    return ("graficos guardados en carpeta")


### GRAFICA DE AC o NMDS ####

def eliminar_unicos_registros (por_tipo):
    """esta funcion elimina los registros unicos que en cada fila sumen uno para el analisis
    de correspondencia
    """
    por_tipo["suma"] = por_tipo.sum(1)
    x = por_tipo[por_tipo["suma"] > 1]
    x = x.drop(['suma'], axis=1)
    por_tipo=por_tipo.drop(["suma"], axis=1)
    return (x)

### definimos la clase ca2 resultante de el objeto ca de ep.
class ca2(ep.ca):
    def site(self):
        return self.siteScores

    def sp(self):
        return self.spScores

def analisis_de_correspondencia (x, annot=True):
    """ realiza un analisis de correspondencia con scaling 1"""
    e = ca2(x, scaling=1)
    coordenadas_sitios = e.site()
    coordenadas_sp = e.sp()
    ax = plt.axes()
    plt.scatter(coordenadas_sitios[0], coordenadas_sitios[1], color="black")
    annotations = coordenadas_sitios.index
    if annot:
        for i, label in enumerate(annotations):
            plt.annotate(label, (coordenadas_sitios[0][i]+0.05, coordenadas_sitios[1][i]+0.05),
                         fontstyle="italic",fontfamily="Arial", fontsize=10)

    plt.scatter(coordenadas_sp[0], coordenadas_sp[1], c="lightgrey", marker="s")
    annotations = coordenadas_sp.index
    for i, label in enumerate(annotations):
        plt.annotate(label, (coordenadas_sp[0][i]+0.05, coordenadas_sp[1][i]+0.05),
                     color="grey",fontfamily="Arial", fontsize=14,
                     )
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    plt.ylabel("CA 2", fontfamily="Arial", fontsize=14)
    plt.xlabel("CA 1", fontfamily="Arial", fontsize=14)
    ##guardamos el grafico obtenido
    plt.savefig(f"analisis_correspondencia.tiff", dpi=300)
    ##obtener el summary de el analisis de correanalisis_correspondenci
    a=e.summary()
    a.to_excel("scalingtipe1-1.xlsx")

# analisis de diversidad automaticos
# method: [‘shannon’ | ‘gini-simpson’ | ‘simpson’ | ‘dominance’ | ‘spRich’ | ‘even’]
# analisis de rarefaccion
# metodos method: [‘rarefy’ | ‘rarecurve’]


def analisis_diversidad_alfa(por_tipo, num_e=True):
    """ funcion que te permite realizar todos los analisis de diversidad alfa y estamacion de rarefaccion de especiesde
    de el data frame, requiere  un dataframe que sea de especies por lugares (columnas)"""
    df = pd.DataFrame()
    names = por_tipo.columns.values
    metodos = ["shannon", "gini-simpson",
               "simpson", "dominance", "spRich", "even"]
    for e in metodos:
        diversidad = ep.diversity(
            por_tipo.transpose(), method=e, breakNA=True, num_equiv=num_e)
        df[e] = diversidad
    df.index = names
    rarefy = ep.rarefy(por_tipo.transpose(), method='rarefy',
                       size=None, breakNA=True)
    df["rarefy"] = rarefy
    df["E(1,0)"] = df["shannon"]/df["spRich"]
    if num_e == True:
        df = df.rename(columns={"spRich": "N(0)", "shannon": "N(1)", "simpson":
                                "N(2)"})
    return df


def grafica_alfa(alfa, indice="shannon", elementos="uno", Type=["CCR","PCR", "SLR"]):
    ax = plt.axes()
    if elementos == "uno":
        index = alfa.index
        values = alfa[indice]
        plt.plot(index, values)
        plt.title(f" Diversidad alfa de {indice}")
    elif elementos == "hill":
        index = alfa.index
        values = ["N(0)", "N(1)", "N(2)"]
        for e in values:
            value = alfa[e]
            plt.plot(index, value, marker=".", linestyle="--")
            value[Type[2]]
            plt.annotate(e, (2.05, value[Type[2]]), fontfamily="Helvetica",
                         fontstyle="italic", fontsize=10)
        #plt.legend(["N(0)","N(1)","N(2)","rarefy", "E(1,0)"])
        #plt.title(f" Diversidad alfa - Numeros de Hill")
        ax.spines['right'].set_visible(False)
        ax.spines['top'].set_visible(False)


###### ANOSIM ####################################
# se tiene que realizar con observaciones mensuales de cada mes
# correspondiente a las colectas sin considerar... falta
# actualizar la base de datos. tambien segun esa informacion se
# va a realizar un NMDS con grupos
def crear_idems (final_bosque_idem):
    """requiere el dato obtenido de final_bosque_idem"""
    a = final_bosque_idem.columns
    grupo = pd.Series(a)
    agroup = [e[0:3] for e in grupo]
    agroup = pd.Series(agroup)
    return (agroup)

def ANOSIM (final_bosque_idem, agroup):
    """requiere el archivo obtenido de final_bosque_idem y el factor de agrupacion obtenida de
    la funcion crear_idems"""
    brayDist = ep.distance(final_bosque_idem. transpose(), method='bray')
    dunesMDS = ep.MDS(brayDist)
    t1 = ep.anosim(brayDist, factor1=agroup)
    #t2 = pd.DataFrame()
    #t2["a"] = t1.summary()
    #t2.to_excel("Anosim.xlsx")
    print(t1.summary())

# MEJORAMOS LOS GRAFICOS DE mds
def NMDS_GRA(final_bosque_idem, criterio_agrupacion=["Pinipini","MLC","Aguanos"]):
    """definimos el grafico de NMDS requiere la informaciond e final_bosque_idem y factor de agrupacion
    obtenida de crear_idems"""
    brayDist = ep.distance(final_bosque_idem. transpose(), method='bray')
    dunesMDS = ep.MDS(brayDist)
    ## buscamos los escores para poder realizar el grafico
    scores = dunesMDS.scores
    scores = pd.DataFrame(scores)
    scores = scores.set_index(pd.Index(agroup))
    ### creamos el grafico
    ax = plt.axes()
    scores1 = scores[scores.index == criterio_agrupacion[0]]
    plt.scatter(scores1[0], scores1[1], marker="x", color="blue")
    poly_sc = a = zip(scores1[0], scores1[1])
    poly_sc = list(poly_sc)
    ax.add_patch(Polygon(poly_sc, color='blue', alpha=0.3))
    plt.annotate(criterio_agrupacion[0], (np.mean(scores1[0]-0.2), np.mean(scores1[1])-0.2),
                 fontfamily="Helvetica", fontsize=10, color="blue", 
                 fontweight="bold")
    scores2 = scores[scores.index == criterio_agrupacion[1]]
    plt.scatter(scores2[0], scores2[1], marker="s", color="green")
    poly_sc = a = zip(scores2[0], scores2[1])
    poly_sc = list(poly_sc)
    ax.add_patch(Polygon(poly_sc, color='forestgreen', alpha=0.2))
    plt.annotate(criterio_agrupacion[1], (np.mean(scores2[0])-0.1, np.mean(scores2[1])-0.1),
                 fontfamily="Helvetiva", fontsize=10, color="green", 
                 fontweight="bold")
    scores3 = scores[scores.index == criterio_agrupacion[2]]
    plt.scatter(scores3[0], scores3[1], marker="o", color="red")
    poly_sc = a = zip(scores3[0], scores3[1])
    poly_sc = list(poly_sc)
    ax.add_patch(Polygon(poly_sc, color='red', alpha=0.2))
    plt.annotate(criterio_agrupacion[2], (np.mean(scores3[0]), np.mean(scores3[1])),
                 fontfamily="Helvetica", fontsize=10, color="red", 
                 fontweight="bold")

    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    plt.ylabel("NMDS 2", fontfamily="Helvetica", fontsize=10)
    plt.xlabel("NMDS 1", fontfamily="Helvetica", fontsize=10)
    plt.savefig(f"NMDS.jpg")

